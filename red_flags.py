import os
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from rule_config import RULES

def build_issue(rule_id, doc_path, paragraph_idx, issue_text, suggestion, rag_engine):
    issue = {
        "rule_id": rule_id,
        "document": os.path.basename(doc_path),
        "section": f"Paragraph {paragraph_idx}" if paragraph_idx else "N/A",
        "issue": issue_text,
        "severity": RULES[rule_id].get("severity", "Medium"),
        "suggestion": suggestion,
        "citation": None,
        "citation_rule": None
    }
    if rag_engine is not None:
        q = RULES[rule_id].get("query_template", "")
        c = rag_engine.get_citation(q)
        if c:
            issue["citation"] = c.get("citation")
            issue["citation_rule"] = c.get("citation_rule")
    return issue

def detect_red_flags(doc_path, rag_engine=None):
    """
    Scans a .docx file for predefined red flags and returns list of issues.
    """
    issues = []
    doc = Document(doc_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip() != ""]

    # iterate paragraphs and apply rules
    for idx, text in enumerate(paragraphs, start=1):
        lower_text = text.lower()

        # Jurisdiction rule
        if "uae federal court" in lower_text or "uae federal courts" in lower_text or "federal court" in lower_text:
            issue = build_issue("JURISDICTION_ADGM", doc_path, idx,
                                "Jurisdiction clause does not specify ADGM",
                                "Update jurisdiction to ADGM Courts.", rag_engine)
            issues.append(issue)

        # You can add more paragraph-level rules here if needed

    # Missing signatory check (end of doc)
    if paragraphs:
        last = paragraphs[-1].lower()
        if "signed by" not in last and "signature" not in last:
            issue = build_issue("MISSING_SIGNATORY", doc_path, len(paragraphs),
                                "Missing signatory section at the end of the document",
                                "Add a proper signatory section with authorized signatures.", rag_engine)
            issues.append(issue)

    return issues

def insert_inline_comments(doc_path, issues, output_path):
    """
    Highlights offending paragraphs (yellow) and appends a RED FLAGS SUMMARY with citations.
    """
    doc = Document(doc_path)
    # Build a paragraph-index map
    para_map = {i+1: p for i, p in enumerate(doc.paragraphs)}

    # Highlight paragraph runs and append text marker to the paragraph
    for issue in issues:
        # find paragraph number from issue['section'] like "Paragraph 3"
        try:
            section = issue.get("section", "")
            if section and section.startswith("Paragraph"):
                num = int(section.split()[1])
            else:
                num = None
        except Exception:
            num = None

        if num and num in para_map:
            paragraph = para_map[num]
            # highlight first run if exists, else create run
            if paragraph.runs:
                run = paragraph.runs[0]
            else:
                run = paragraph.add_run()
            try:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            except Exception:
                # in case highlight not supported, append marker
                pass
            # append a short marker
            paragraph.add_run(f"  [⚠ {issue['severity']} - {issue['issue']}]")

    # Append summary section with suggestions and citations
    if issues:
        doc.add_page_break()
        doc.add_paragraph("=== RED FLAGS SUMMARY ===")
        for issue in issues:
            p = doc.add_paragraph()
            p.add_run(f"⚠ {issue['severity']} - {issue['issue']}").bold = True
            p.add_run(f"\nSuggestion: {issue['suggestion']}")
            if issue.get("citation"):
                c = doc.add_paragraph()
                c.add_run(f"Citation (source: {issue.get('citation_rule')}): ").italic = True
                c.add_run(issue.get("citation")[:1000] + "...")

    doc.save(output_path)